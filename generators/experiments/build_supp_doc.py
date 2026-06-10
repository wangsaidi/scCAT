"""
build_supp_doc.py — Compile all 14 supplementary figures (S1-S14) into a
single Word document with a cover page, a Contents page (TOC) and one page
per figure (heading + embedded PNG + detailed English caption).

Output: cluade-version/Supplementary_Figures_S1-S14.docx
"""

from __future__ import annotations
from pathlib import Path

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

HERE = Path(__file__).resolve().parent
BASE = HERE.parent.parent
FIG_DIR = BASE / "resource" / "figures"
OUT_PATH = BASE / "Supplementary_Figures_S1-S14.docx"

# ─────────────────────────────────────────────────────────────────────────────
# Style
# ─────────────────────────────────────────────────────────────────────────────

EN_FONT = "Times New Roman"
CN_FONT = "Microsoft YaHei"


def _set_run_fonts(run, cn=CN_FONT, en=EN_FONT, size_pt=10.5):
    run.font.name = en
    run.font.size = Pt(size_pt)
    r_pr = run._element.get_or_add_rPr()
    rFonts = r_pr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        r_pr.append(rFonts)
    rFonts.set(qn("w:eastAsia"), cn)


def add_para(doc, text, bold=False, italic=False, size_pt=10.5,
             align=None, first_line_indent_cm=0, space_after_pt=4,
             color_hex=None):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.first_line_indent = Cm(first_line_indent_cm)
    pf.space_after = Pt(space_after_pt)
    if align == "center":
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif align == "right":
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = p.add_run(text)
    r.bold = bold
    r.italic = italic
    if color_hex:
        r.font.color.rgb = RGBColor.from_string(color_hex)
    _set_run_fonts(r, size_pt=size_pt)
    return p, r


def add_title(doc, text, level=1, align="center"):
    sizes = {0: 20, 1: 16, 2: 13, 3: 11}
    size_pt = sizes.get(level, 11)
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(8)
    if align == "center":
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.bold = True
    _set_run_fonts(r, size_pt=size_pt)
    return p


def add_page_break(doc):
    p = doc.add_paragraph()
    p.add_run().add_break(WD_BREAK.PAGE)


def add_caption(doc, fig_id, fig_title, body):
    """Caption block: bold "Supplementary Figure SN | Title." prefix +
    descriptive body text (panel-by-panel)."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.first_line_indent = Cm(0)

    # Bold prefix
    r_prefix = p.add_run(f"Supplementary Figure {fig_id} | {fig_title}. ")
    r_prefix.bold = True
    _set_run_fonts(r_prefix, size_pt=10)

    # Body in regular weight
    r_body = p.add_run(body)
    _set_run_fonts(r_body, size_pt=10)


def add_figure_image(doc, fig_path: Path, width_cm: float = 16.0):
    """Insert a centered image, scaled to width_cm. Caller controls
    the surrounding paragraph layout."""
    if not fig_path.exists():
        # Placeholder if image missing
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(f"[image missing: {fig_path.name}]")
        r.italic = True
        r.font.color.rgb = RGBColor(0xCC, 0x33, 0x33)
        _set_run_fonts(r, size_pt=10)
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run()
    run.add_picture(str(fig_path), width=Cm(width_cm))


# ─────────────────────────────────────────────────────────────────────────────
# Figure catalogue — file basename, short title, detailed legend
# ─────────────────────────────────────────────────────────────────────────────

FIGURES = [
    {
        "id":       "S1",
        "file":     "FigS1_sim1",
        "short":    "Full results on Simulated Dataset 1",
        "title":    "Full 12-method results on Simulated Dataset 1 "
                    "(3 batches × 4 balanced cell types)",
        "legend": (
            "(a) UMAP visualisations of the integrated embeddings from all "
            "twelve methods, coloured by batch. (b) The same UMAPs coloured "
            "by cell type. (c) Bottom-left: LISI scatter showing cLISI "
            "(cell-type purity, higher is better) vs iLISI (batch mixing, "
            "higher is better); each dot is one method, scCAT highlighted "
            "in green. Bottom-middle: kBET rejection rate (lower = better) "
            "as horizontal lollipops. Bottom-right (panel d): Evaluation "
            "Summary heatmap, methods × six canonical metrics, colour-coded "
            "by per-column min-max normalised score. This balanced simulation "
            "is the simplest test scenario: all methods perform well; scCAT "
            "ranks at or near the top on every metric while showing no "
            "evidence of overcorrection."
        ),
    },
    {
        "id":       "S2",
        "file":     "FigS2_sim2",
        "short":    "Partial-sharing test on Simulated Dataset 2",
        "title":    "Full 12-method results on Simulated Dataset 2 "
                    "(Type 4 batch-specific to Batch 3)",
        "legend": (
            "Same layout as Suppl Fig. S1. Type 4 cells are present only in "
            "Batch 3; on the cell-type UMAPs (b) batch-specific cells are "
            "outlined in black. fastMNN and Scanorama begin to absorb the "
            "Type 4 cells into the other clusters (visible as the loss of "
            "Group4 as an independent cluster); scCAT preserves Group4 as a "
            "distinct cluster, foreshadowing the same pattern observed on "
            "the larger Sim 4 dataset (Fig. 2) and on the real HDC dataset "
            "(Fig. 3)."
        ),
    },
    {
        "id":       "S3",
        "file":     "FigS3_sim3",
        "short":    "Full 12-method view of Sim 3 (main Fig. 2 shows 4)",
        "title":    "Full 12-method results on Simulated Dataset 3 "
                    "(10 batches × 10 fully shared cell types)",
        "legend": (
            "Main Fig. 2 presents only four representative methods "
            "(Scanorama, fastMNN, SPDR, scCAT). This supplementary figure "
            "shows the full twelve-method panel under the same layout as "
            "Suppl Fig. S1. All ten cell types are shared across all ten "
            "batches; the principal differences between methods appear in "
            "the degree of batch mixing rather than in cell-type structure. "
            "scCAT and fastMNN both achieve ARI = 1.00 on this benchmark; "
            "their iLISI and ASW_batch_mixing differ marginally."
        ),
    },
    {
        "id":       "S4",
        "file":     "FigS4_sim4",
        "short":    "Full 12-method view of Sim 4 (main Fig. 2 shows 4)",
        "title":    "Full 12-method results on Simulated Dataset 4 "
                    "(Group1/3/5/7 batch-specific)",
        "legend": (
            "Main Fig. 2 presents only four representative methods on Sim 4; "
            "this supplementary figure shows the full twelve-method panel. "
            "Group1, Group3, Group5 and Group7 exist only in Batch 2, 4, 6 "
            "and 8 respectively; these cells are outlined in black on the "
            "cell-type UMAPs (b). fastMNN and INSCT merge all four "
            "batch-specific cell types into a single cluster — this is the "
            "failure mode that drives high OCI in main Fig. 2d. Scanorama "
            "and DESC partially preserve the batch-specific structure but "
            "lose batch mixing in the shared clusters. Only scCAT achieves "
            "concurrent high batch mixing and complete preservation of all "
            "four batch-specific cell types."
        ),
    },
    {
        "id":       "S5",
        "file":     "FigS5_scmixology",
        "short":    "Full 12-method view of Sc_mixology",
        "title":    "Full 12-method results on Sc_mixology "
                    "(cross-platform integration)",
        "legend": (
            "1,401 cells comprising three lung-cancer cell lines (HCC827, "
            "H1975, H2228) profiled on three sequencing platforms (10x "
            "Chromium, CEL-seq2, Drop-seq). All three cell lines are shared "
            "across all three platforms — this is a pure technical-batch "
            "integration task. scCAT and SPDR achieve the cleanest cell-line "
            "separation with concurrent platform mixing; DESC and scBCN "
            "show residual cell-type confusion (visible as colour mixing "
            "in cell-type UMAPs). The multi-seed (3-seed) IB ranked bar "
            "for Sc_mixology is in main Fig. 5g, confirming that scCAT's "
            "advantage is reproducible across random seeds."
        ),
    },
    {
        "id":       "S6",
        "file":     "FigS6_pbmc",
        "short":    "Full 12-method view of PBMC",
        "title":    "Full 12-method results on PBMC "
                    "(control vs IFN-β stimulation)",
        "legend": (
            "13,576 peripheral blood mononuclear cells across two conditions "
            "(control and IFN-β stimulated) and eight major immune cell types "
            "(CD4 T, CD8 T, B, NK, CD14+ monocytes, CD16+ monocytes, dendritic "
            "cells, megakaryocytes). This is a condition-associated integration "
            "task — the two 'batches' differ both technically and biologically. "
            "scCAT achieves the most favorable balance between condition mixing "
            "and cell-type preservation; DeepBID and DESC retain visible "
            "condition separation. The biological-readout panel (main Fig. 5f) "
            "confirms that scCAT preserves the IFN-response signature in the "
            "stimulated condition rather than treating it as a removable "
            "batch effect. The per-cell-type IFN-β DE concordance across "
            "four methods (scCAT, Harmony, Scanorama, INSCT) is quantified "
            "in main Fig. 5h."
        ),
    },
    {
        "id":       "S7",
        "file":     "FigS7_pancreas",
        "short":    "Real-data validation on Human Pancreas",
        "title":    "Full 12-method results on the Human Pancreas dataset "
                    "(9 batches × 14 cell types)",
        "legend": (
            "Multi-protocol integration of pancreatic islet cells across nine "
            "studies. The dataset is compositionally heterogeneous: rare cell "
            "types (e.g. mast, schwann, epsilon cells) are present only in a "
            "subset of batches. scCAT preserves all 14 cell types as distinct "
            "clusters while integrating the four major endocrine populations "
            "(alpha, beta, delta, gamma) across protocols."
        ),
    },
    {
        "id":       "S8",
        "file":     "FigS8_immune",
        "short":    "Real-data validation on Human Immune",
        "title":    "Full 12-method results on the Human Immune dataset "
                    "(10 donors / studies × 16 immune cell types)",
        "legend": (
            "Multi-donor integration of 33,506 immune cells across ten "
            "donors and studies. This is the largest atlas-level real "
            "dataset in our benchmark suite besides Mouse Lung. Rare cell "
            "types (CD10+ B, Erythrocytes, NK_dim, CD8+ TEM_KLRG1) appear "
            "in only a subset of donors — scCAT preserves them as distinct "
            "clusters; competing methods absorb them into neighbouring "
            "abundant cell types."
        ),
    },
    {
        "id":       "S9",
        "file":     "FigS9_gut",
        "short":    "Real-data validation on Gut (condition)",
        "title":    "Full 12-method results on the Gut dataset "
                    "(4 conditions × 8 epithelial cell types)",
        "legend": (
            "9,842 mouse intestinal epithelial cells across four conditions "
            "(Control, Salmonella infection, H. polygyrus infection, and a "
            "second control). Condition-induced cell states (e.g. expanded "
            "tuft cells under H. polygyrus, expanded enterocytes under "
            "Salmonella) are biologically meaningful and should not be "
            "removed by integration. scCAT preserves these condition-specific "
            "expansions while integrating the shared progenitor populations "
            "across conditions."
        ),
    },
    {
        "id":       "S10",
        "file":     "FigS10_sensitivity",
        "short":    "Parameter sensitivity of scCAT on HDC",
        "title":    "Parameter sensitivity analysis of scCAT on HDC",
        "legend": (
            "Five core scCAT parameters — min_c_pos (positive-confidence "
            "threshold), m₀ (base margin), μ_rare (batch-specific protection "
            "weight), γ (confidence-weight scale) and k (same-batch KNN) — "
            "were swept across 4–5 values each (24 configurations total). "
            "Each row corresponds to one parameter; the three columns report "
            "OCI (↓), BSRS (↑) and Integration Balance (↑) computed on the "
            "HDC dataset. Black stars mark the default value used elsewhere "
            "in the manuscript. The IB range within any single parameter "
            "sweep is below 0.05, demonstrating that scCAT is robust across "
            "the tested ranges — readers reproducing the method are unlikely "
            "to encounter strong sensitivity to minor parameter perturbations. "
            "Raw values for all 24 configurations are tabulated in "
            "Supplementary Table S13."
        ),
    },
    {
        "id":       "S11",
        "file":     "FigS11_runtime",
        "short":    "Runtime and peak memory of scCAT",
        "title":    "Runtime and peak memory of scCAT (CPU)",
        "legend": (
            "(a) Wall-clock training time of scCAT on four benchmark "
            "datasets (HDC, Sc_mixology, PBMC, Sim 4). (b) Peak Python "
            "memory (RSS measured by psutil + tracemalloc). (c) Log-log "
            "scaling of runtime against cell count; minor ticks within each "
            "decade are shown for clarity, and dataset labels are routed "
            "with adjustText to avoid overlapping the data points. (d) "
            "Numerical summary table. All measurements were performed on a "
            "single CPU configuration (Intel x86_64, no GPU). For the "
            "10x 73k PBMC down-sampling scaling curve (main Fig. 6f, g) "
            "we use the same protocol with subsampled cell counts; "
            "cross-method wall-clock comparison is out of scope and is "
            "discussed in Methods §4.7. The Wilcoxon BH-corrected p-value "
            "summary (scCAT vs all methods) is in main Fig. 6h."
        ),
    },
    {
        "id":       "S12",
        "file":     "FigS12_ext_ablation",
        "short":    "Extended ablation on Sim 4 + HDC",
        "title":    "Module ablations on Sim 4 and HDC — extended view",
        "legend": (
            "Each of five scCAT variants (full, −Conf, −Filter, "
            "−AdaptMargin, −BSP) is evaluated on Sim 4 (top row) and HDC "
            "(bottom row), across four metrics: OCI (↓), BSRS (↑), "
            "Integration Balance (↑) and clustering ARI (↑). scCAT-full "
            "is marked with a green star. On Sim 4, removing the "
            "batch-specific protection (BSP) increases OCI from 0.01 to "
            "0.26 and removing the confidence-weighting module (−Conf) "
            "increases OCI from 0.01 to 0.75 — the same pattern reported "
            "for IB in main Fig. 4 b–d. On HDC, all module effects are "
            "compressed into a narrow IB range because of small-sample "
            "noise (n = 569 cells), but the relative ordering remains "
            "consistent: −Conf and −Filter together remove the most signal. "
            "Numerical source data are tabulated in Supplementary Table S7."
        ),
    },
    {
        "id":       "S13",
        "file":     "FigS13_lung_full",
        "short":    "Full 12-method Mouse Lung view (complements Fig. 6c)",
        "title":    "Full 12-method results on the Mouse Lung atlas "
                    "(32,472 cells × 16 batches × 17 cell types)",
        "legend": (
            "Top (a): UMAP grid coloured by batch — 16 batches indicated by "
            "the colour palette to the right. Bottom (b): the same UMAPs "
            "coloured by cell type — 17 lung cell types. This figure "
            "complements the per-batch local-mixing heatmap in main "
            "Fig. 6c, which condenses the twelve UMAPs in (a) into a "
            "single 16 × 12 numerical heatmap; together they provide both "
            "the visual and the quantitative view of batch mixing in this "
            "atlas-level dataset. scCAT and SPDR achieve the cleanest "
            "cell-type separation while maintaining batch mixing across "
            "all 16 batches; Scanorama achieves very high batch mixing "
            "but at the cost of cell-type structure (visible as cell-type "
            "colours bleeding into each other)."
        ),
    },
    {
        "id":       "S14",
        "file":     "FigS14_scIB_robustness",
        "short":    "scCAT's ranking is robust to the aggregation choice",
        "title":    "scCAT's standing does not depend on the Integration "
                    "Balance (IB) aggregation",
        "legend": (
            "Methods were re-ranked on scCAT's nine evaluation datasets "
            "(four simulations + five real datasets; random seeds averaged "
            "within each dataset) using the community-standard scIB "
            "aggregation (Luecken et al. 2022), scIB_overall = "
            "0.4·s_batch + 0.6·s_bio, computed from the same component "
            "metrics as IB but with the standard linear weighting instead "
            "of our geometric mean (s_batch = mean of iLISI, 1 − kBET, "
            "ASW_batch_mixing; s_bio = mean of ARI, NMI, ASW_celltype, "
            "cLISI_purity). The analysis is restricted to the nine methods "
            "with complete coverage across these datasets. "
            "(a) Mean community-standard scIB overall score per method "
            "(higher = better). scCAT (0.716, bold) attains the best score "
            "among all unsupervised methods and the second-best overall, "
            "behind only the semi-supervised scANVI (0.734; † uses "
            "ground-truth cell-type labels). Bars for scCAT and scANVI are "
            "highlighted and values are printed at the bar ends. "
            "(b) Per-method average rank under our IB (x-axis) versus under "
            "the community-standard scIB (y-axis); both axes are oriented "
            "so that rank 1 (best) is at the top-right and the dashed line "
            "is the identity. The two rankings agree almost perfectly "
            "(Spearman ρ = 0.98): every method lies on or next to the "
            "diagonal, and scANVI and scCAT occupy the top two ranks under "
            "both aggregations. The method ordering is therefore essentially "
            "invariant to the choice of aggregation, showing that scCAT's "
            "standing reflects the underlying component metrics rather than "
            "the IB formula. A dataset-level Friedman test across the nine "
            "methods is highly significant (χ² = 43.3, p = 7.7×10⁻⁷, "
            "k = 9, N = 9; Nemenyi critical difference = 4.00). Numerical "
            "values are in Supplementary Table S18."
        ),
    },
]


# ─────────────────────────────────────────────────────────────────────────────
# Document build
# ─────────────────────────────────────────────────────────────────────────────

def build_supp_doc():
    doc = Document()

    # Page setup — A4 portrait, 2.5 cm margins (same as main manuscript)
    for section in doc.sections:
        section.top_margin    = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.5)

    # ──────────────────────────────────────────────────────────────────
    # Page 1 — title page
    # ──────────────────────────────────────────────────────────────────
    add_title(
        doc,
        "Supplementary Information",
        level=0,
    )
    add_title(
        doc,
        "scCAT preserves batch-specific cell states through "
        "confidence-weighted triplet learning in single-cell data integration",
        level=2,
    )
    add_para(doc,
             "Supplementary Figures S1–S14",
             align="center", bold=True, size_pt=12, space_after_pt=6)
    add_para(doc,
             "(Supplementary Tables S1–S19 are provided separately as "
             "Supplementary_Tables_S1-S19.xlsx)",
             align="center", italic=True, size_pt=10, space_after_pt=12,
             color_hex="666666")

    # Authors placeholder
    add_para(doc,
             "[Author list — paste from main manuscript]",
             align="center", italic=True, size_pt=10,
             color_hex="888888")
    add_para(doc,
             "[Affiliations and corresponding author email]",
             align="center", italic=True, size_pt=10,
             color_hex="888888")

    add_page_break(doc)

    # ──────────────────────────────────────────────────────────────────
    # Page 2 — Contents (TOC)
    # ──────────────────────────────────────────────────────────────────
    add_title(doc, "Contents", level=1, align="center")
    add_para(doc,
             "The following 14 supplementary figures accompany the main "
             "manuscript. Each figure is shown at full size with a detailed "
             "panel-by-panel legend on the page where the figure appears.",
             italic=True, size_pt=10, space_after_pt=10)

    for fig in FIGURES:
        p, _ = add_para(doc, "", space_after_pt=2)
        # "Supplementary Figure SN" — bold
        r_id = p.add_run(f"Supplementary Figure {fig['id']}")
        r_id.bold = True
        _set_run_fonts(r_id, size_pt=10.5)
        # tab + short description
        r_desc = p.add_run(f" — {fig['short']}")
        _set_run_fonts(r_desc, size_pt=10.5)

    add_page_break(doc)

    # ──────────────────────────────────────────────────────────────────
    # Pages 3+ — one figure per page
    # ──────────────────────────────────────────────────────────────────
    for i, fig in enumerate(FIGURES):
        fig_path = FIG_DIR / f"{fig['file']}.png"

        # Page heading
        add_title(doc, f"Supplementary Figure {fig['id']}",
                  level=1, align="center")

        # Image (centred, 16 cm wide)
        add_figure_image(doc, fig_path, width_cm=16.0)

        # Caption (justified would be ideal; left-aligned by default)
        add_caption(doc, fig['id'], fig['title'], fig['legend'])

        # Page break after each figure (except the last)
        if i < len(FIGURES) - 1:
            add_page_break(doc)

    # ──────────────────────────────────────────────────────────────────
    # Save
    # ──────────────────────────────────────────────────────────────────
    doc.save(OUT_PATH)
    print(f"Saved: {OUT_PATH}")
    return OUT_PATH


if __name__ == "__main__":
    build_supp_doc()
